# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START, WD_SECTION_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from pathlib import Path

OUTPUT = Path("九江学院2023—2025年临床执业医师资格考试综合分析报告_正式版.docx")

BODY_FONT = "仿宋_GB2312"
HEADING_FONT = "黑体"
KAI_FONT = "楷体_GB2312"
TABLE_FONT = "宋体"
EN_FONT = "Times New Roman"


def set_run_font(run, chinese=BODY_FONT, size=16, bold=None, color=None):
    run.font.name = EN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_cell_text(cell, text, size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(str(text))
    set_run_font(r, chinese=TABLE_FONT, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill="D9E2F3"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_cm=16.0):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_cm * 567)))
    tbl_w.set(qn("w:type"), "dxa")


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, chinese=TABLE_FONT, size=12, bold=True)
    return p


def add_table(doc, headers, rows, widths=None, font_size=10.5, header_fill="D9E2F3", first_col_left=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_width(table)
    hdr = table.rows[0]
    repeat_header(hdr)
    prevent_row_split(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, size=font_size, bold=True)
        shade_cell(hdr.cells[i], header_fill)
        set_cell_margins(hdr.cells[i])
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if first_col_left and i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(row.cells[i], val, size=font_size, bold=False, align=align)
            set_cell_margins(row.cells[i])
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, chinese=TABLE_FONT, size=9.5)
    return p


def add_body(doc, text, indent=True, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(32) if indent else Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, chinese=BODY_FONT, size=16, bold=bold)
    return p


def add_item(doc, number, title, text, responsibility=None):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{number}. {title}")
    set_run_font(r1, chinese=BODY_FONT, size=16, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, chinese=BODY_FONT, size=16)
    if responsibility:
        r3 = p.add_run(f"（牵头部门：{responsibility}）")
        set_run_font(r3, chinese=KAI_FONT, size=15)
    return p


def add_heading_cn(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5 if level == 1 else 2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(30 if level == 1 else 28)
    for r in p.runs:
        set_run_font(r, chinese=HEADING_FONT if level <= 2 else KAI_FONT,
                     size=20 if level == 1 else 16, bold=True)
    return p


def add_page_number(section, start=1):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    set_run_font(run, chinese=TABLE_FONT, size=10.5)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def add_toc_field(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在打开文档时自动更新"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])
    set_run_font(run, chinese=TABLE_FONT, size=12)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(16)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(28)
    normal.paragraph_format.space_after = Pt(0)

    for level in (1, 2, 3):
        style = styles[f"Heading {level}"]
        style.font.name = EN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT if level <= 2 else KAI_FONT)
        style.font.size = Pt(20 if level == 1 else 16)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True


def set_update_fields(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def all_document_text(doc):
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


doc = Document()
configure_styles(doc)
set_update_fields(doc)

# Page setup
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

# Document properties
doc.core_properties.title = "九江学院2023—2025年临床执业医师资格考试综合分析报告"
doc.core_properties.subject = "临床医学人才培养质量分析与改进"
doc.core_properties.author = "九江学院医学部"
doc.core_properties.keywords = "医师资格考试；临床医学；人才培养质量；持续改进"

# Cover
cover_sec = doc.sections[0]
cover_sec.vertical_alignment = WD_SECTION_VERTICAL_ALIGNMENT.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("九江学院")
set_run_font(r, chinese=HEADING_FONT, size=24, bold=True)
p.paragraph_format.space_after = Pt(28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("2023—2025年临床执业医师资格考试\n综合分析报告")
set_run_font(r, chinese="方正小标宋简体", size=28, bold=False)
p.paragraph_format.line_spacing = 1.35
p.paragraph_format.space_after = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("（全日制应届本科考生）")
set_run_font(r, chinese=KAI_FONT, size=17)
p.paragraph_format.space_after = Pt(90)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("医学部")
set_run_font(r, chinese=BODY_FONT, size=18)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("2026年6月")
set_run_font(r, chinese=BODY_FONT, size=18)

# Contents section
contents_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
contents_sec.vertical_alignment = WD_SECTION_VERTICAL_ALIGNMENT.TOP
contents_sec.top_margin = Cm(2.6)
contents_sec.bottom_margin = Cm(2.4)
contents_sec.left_margin = Cm(2.8)
contents_sec.right_margin = Cm(2.6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("目  录")
set_run_font(r, chinese=HEADING_FONT, size=20, bold=True)
p.paragraph_format.space_after = Pt(18)
add_toc_field(doc)

# Body section
body_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
body_sec.top_margin = Cm(2.6)
body_sec.bottom_margin = Cm(2.4)
body_sec.left_margin = Cm(2.8)
body_sec.right_margin = Cm(2.6)
add_page_number(body_sec, 1)

# 一、前言
add_heading_cn(doc, "一、前言", 1)
add_body(doc, "医师资格考试是检验医学毕业生达到行业准入要求的重要外部评价指标，也是学校开展人才培养质量监测和持续改进的重要依据。本报告依据国家医学考试中心提供的2023—2025年《医师资格考试成绩分析报告（临床执业医师）》，结合《中国本科医学教育标准——临床医学专业（2022版）》相关要求，对九江学院近三年临床执业医师资格考试情况进行系统分析，明确优势与不足，提出下一阶段改进措施，为学校教学指导委员会和校长办公会研究医学教育教学工作提供参考。")
add_body(doc, "九江学院医学教育历史可追溯至1901年创办的但福德护士学校，2000年开设临床医学专业本科教育，2015年首次通过教育部临床医学专业认证，临床医学专业现为江西省一流本科专业建设点。2025年6月，教育部临床医学专业认证工作委员会专家组到校开展现场考察指导，对学校医学教育基础、办学传统和组织体系给予肯定。学校应以新一轮专业认证和医学教育“十五五”规划编制为契机，进一步完善人才培养体系，持续提高毕业生岗位胜任力。")
add_body(doc, "本报告以国家医学考试中心定义的全日制应届本科考生为主要分析对象，即毕业后一年参加医师资格考试的全日制本科毕业生。2023年、2024年和2025年报告所对应的毕业年份分别为2022年、2023年和2024年。全体考生及非应届考生数据作为毕业后跟踪服务和学校考试总体情况的补充资料。具体数据口径见附件。")

# 二、总体情况
add_heading_cn(doc, "二、近三年考试总体情况", 1)
add_heading_cn(doc, "2.1 实践技能考试", 2)
add_caption(doc, "表1  2023—2025年实践技能考试情况（全日制应届本科考生）")
add_table(doc,
          ["年份", "本校实考人数", "本校通过率", "全国通过率", "本校与全国差值", "百分等级"],
          [
              ["2023年", "162", "91.36%", "90.59%", "+0.77个百分点", "—"],
              ["2024年", "255", "90.59%", "91.27%", "-0.68个百分点", "D"],
              ["2025年", "250", "90.40%", "91.73%", "-1.33个百分点", "D"],
          ], widths=[2.0, 2.7, 2.7, 2.7, 3.1, 1.5], font_size=10.5)
add_source_note(doc, "数据来源：国家医学考试中心2023—2025年《医师资格考试成绩分析报告（临床执业医师）》。")
add_body(doc, "本校全日制应届本科考生实践技能考试通过率近三年稳定在90%以上。2023年高于全国0.77个百分点，2024年和2025年分别低于全国0.68和1.33个百分点，连续两年百分等级为D。实践技能平均分由2023年的72.73分上升至2024年的73.44分，2025年回落至71.90分。学校应进一步提高技能训练的规范性、覆盖度和持续性，重点加强基本操作、综合技能、临床思维和医患沟通的协同训练。")

add_heading_cn(doc, "2.2 医学综合考试", 2)
add_caption(doc, "表2  2023—2025年医学综合考试通过情况（全日制应届本科考生）")
add_table(doc,
          ["年份", "本校实考人数", "本校通过率", "全国通过率", "本校与全国差值", "百分等级"],
          [
              ["2023年", "146", "70.55%", "74.24%", "-3.69个百分点", "—"],
              ["2024年", "229", "82.53%", "82.69%", "-0.16个百分点", "C"],
              ["2025年", "224", "80.80%", "82.32%", "-1.52个百分点", "D"],
          ], widths=[2.0, 2.7, 2.7, 2.7, 3.1, 1.5], font_size=10.5)
add_caption(doc, "表3  2023—2025年医学综合考试平均分比较")
add_table(doc,
          ["年份", "本校平均分", "全国平均分", "本校与全国差值"],
          [
              ["2023年", "379.83", "385.80", "-5.97分"],
              ["2024年", "407.34", "409.71", "-2.37分"],
              ["2025年", "396.68", "405.71", "-9.03分"],
          ], widths=[3.0, 4.0, 4.0, 4.0], font_size=10.5)
add_source_note(doc, "数据来源：国家医学考试中心2023—2025年《医师资格考试成绩分析报告（临床执业医师）》。")
add_body(doc, "2023年本校医学综合考试通过率低于全国3.69个百分点。2024年通过率提高至82.53%，与全国平均水平基本持平，百分等级达到C；2025年通过率为80.80%，低于全国1.52个百分点，百分等级回落至D。医学综合考试平均分与全国的差距由2023年的5.97分缩小至2024年的2.37分，2025年扩大至9.03分。结果显示，本校医学综合考试具备达到全国平均水平的基础，但年度表现的稳定性仍需加强。")

add_heading_cn(doc, "2.3 总通过率", 2)
add_caption(doc, "表4  2023—2025年总通过率（全日制应届本科考生）")
add_table(doc,
          ["年份", "本校总通过率", "全国总通过率", "本校与全国差值", "百分等级"],
          [
              ["2023年", "63.58%", "66.59%", "-3.01个百分点", "—"],
              ["2024年", "74.12%", "74.67%", "-0.55个百分点", "D"],
              ["2025年", "72.40%", "74.78%", "-2.38个百分点", "D"],
          ], widths=[2.4, 3.0, 3.0, 3.5, 2.0], font_size=10.5)
add_body(doc, "近三年本校总通过率分别为63.58%、74.12%和72.40%，均低于同期全国平均水平。2024年与全国差距缩小至0.55个百分点，2025年扩大至2.38个百分点；2024年和2025年百分等级均为D。学校下一阶段应以稳定达到全国平均水平为基本目标，同时关注百分等级、平均分和连续年度表现。")

add_heading_cn(doc, "2.4 2025年外部比较", 2)
add_caption(doc, "表5  2025年医学综合考试通过率及平均分比较")
add_table(doc,
          ["考生群体", "通过率", "平均分"],
          [
              ["本校", "80.80%", "396.68"],
              ["全国", "82.32%", "405.71"],
              ["辖区内学校", "86.02%", "408.58"],
              ["“双一流”高校", "90.48%", "421.01"],
              ["原211高校", "89.49%", "418.40"],
              ["原985高校", "94.53%", "429.13"],
              ["通过率最高学校", "100.00%", "437.83"],
          ], widths=[6.0, 4.5, 4.5], font_size=10.5, first_col_left=True)
add_body(doc, "2025年本校医学综合考试通过率低于辖区内学校5.22个百分点，平均分低于辖区内学校11.90分；与“双一流”、原211、原985高校相比，差距进一步扩大。学校应将辖区内学校平均水平作为近期追赶目标，将高水平医学院校作为课程建设、考核评价和临床教学改革的重要参照。")

add_heading_cn(doc, "2.5 影响总通过率的主要环节", 2)
add_body(doc, "2023—2025年，三届全日制应届本科考生共有667人参加实践技能考试，605人通过实践技能考试，599人参加医学综合考试，473人最终通过医学综合考试。在194名未最终通过的考生中，62人未通过实践技能考试，6人通过实践技能考试后未参加医学综合考试，126人未通过医学综合考试。未最终通过者中约三分之二集中在医学综合考试环节，医学综合考试是当前提升总通过率的主要着力点。")

# 三、模块与知识点
add_heading_cn(doc, "三、模块与专业知识点掌握情况", 1)
add_heading_cn(doc, "3.1 各模块整体情况", 2)
add_caption(doc, "表6  2023—2025年各模块掌握率及与全国差值")
add_table(doc,
          ["模块", "2023年本校", "2023年差值", "2024年本校", "2024年差值", "2025年本校", "2025年差值"],
          [
              ["基础医学", "57.72%", "-1.16", "58.71%", "-1.05", "60.85%", "-2.21"],
              ["医学人文", "65.50%", "-1.96", "58.72%", "-0.70", "66.26%", "-1.47"],
              ["临床医学", "64.00%", "-0.76", "71.03%", "-0.05", "67.52%", "-1.38"],
              ["预防医学", "63.39%", "-2.19", "61.04%", "-2.35", "61.09%", "-1.89"],
              ["中医学基础", "—", "—", "63.54%", "-3.64", "55.28%", "-0.55"],
          ], widths=[2.6, 2.1, 1.8, 2.1, 1.8, 2.1, 1.8], font_size=9.5)
add_source_note(doc, "注：差值为本校掌握率减全国掌握率，单位为百分点。2023年考试未设置中医学基础模块。")
add_body(doc, "2023年四个模块、2024年和2025年五个模块的本校掌握率均低于同期全国平均水平。临床医学模块相对表现较好，2024年与全国仅相差0.05个百分点，2025年差距扩大至1.38个百分点。基础医学和预防医学连续三年低于全国平均水平，是需要持续加强的两个模块。医学人文模块年度波动较大，应结合课程考核和年度试题结构开展专题分析。")

add_heading_cn(doc, "3.2 基础医学模块", 2)
add_caption(doc, "表7  基础医学各学科或专业知识点掌握率及与全国差值")
add_table(doc,
          ["学科或知识点", "2023年本校", "2023年差值", "2024年本校", "2024年差值", "2025年本校", "2025年差值"],
          [
              ["病理生理学", "50.38%", "-2.37", "71.45%", "-1.33", "73.83%", "-2.24"],
              ["病理学", "65.71%", "-1.29", "69.03%", "-1.26", "68.72%", "-0.74"],
              ["解剖学", "51.44%", "-3.56", "24.18%", "-4.12", "44.98%", "-5.26"],
              ["生理学", "62.46%", "+0.92", "59.89%", "-0.29", "65.08%", "-0.59"],
              ["生物化学", "59.60%", "-1.40", "65.59%", "-0.72", "60.27%", "-1.89"],
              ["药理学", "48.25%", "-2.38", "53.46%", "-0.81", "61.35%", "-2.98"],
              ["医学免疫学", "64.00%", "+0.33", "52.91%", "+0.46", "40.48%", "-5.63"],
              ["医学微生物学", "51.33%", "+0.66", "61.57%", "-0.05", "57.96%", "-1.24"],
          ], widths=[2.7, 2.0, 1.8, 2.0, 1.8, 2.0, 1.8], font_size=9.0, first_col_left=True)
add_body(doc, "（1）解剖学连续三年低于全国平均水平，与全国的差距分别为3.56、4.12和5.26个百分点，表现为持续性相对薄弱。2025年原始掌握率较2024年有所回升，但与全国差距继续扩大，仍应作为基础医学重点改进学科。")
add_body(doc, "（2）医学免疫学在2023年和2024年分别高于全国0.33和0.46个百分点，2025年低于全国5.63个百分点，年度变化较为明显。学校应组织课程组开展专项分析，重点检查课程目标、教学内容、学时安排、师资变化、校内考核和学生学习情况，形成有针对性的改进方案。")
add_body(doc, "（3）药理学原始掌握率逐年提高，但2025年与全国差距扩大至2.98个百分点；病理生理学2025年低于全国2.24个百分点。上述学科应进一步加强临床情境中的知识应用和综合分析训练。")
add_body(doc, "（4）生理学和医学微生物学总体接近全国平均水平。生理学2023年高于全国0.92个百分点，2024年和2025年分别低于全国0.29和0.59个百分点；医学微生物学三年差值分别为+0.66、-0.05和-1.24个百分点。应巩固现有基础并关注年度变化。")

add_heading_cn(doc, "3.3 临床医学模块", 2)
add_caption(doc, "表8  临床医学主要学科或专业知识点掌握率及与全国差值")
add_table(doc,
          ["学科或知识点", "2023年本校", "2023年差值", "2024年本校", "2024年差值", "2025年本校", "2025年差值"],
          [
              ["内科学", "64.36%", "-0.67", "72.79%", "+0.33", "68.93%", "-1.50"],
              ["神经内科", "—", "—", "62.34%", "-1.67", "64.79%", "-4.14"],
              ["肾脏内科", "—", "—", "70.21%", "+0.43", "58.79%", "-2.20"],
              ["血液内科", "—", "—", "60.40%", "+1.80", "52.85%", "-1.27"],
              ["风湿免疫", "—", "—", "60.39%", "+0.52", "54.08%", "-1.87"],
              ["外科学", "65.01%", "-0.71", "69.82%", "-0.21", "67.21%", "-1.92"],
              ["普通外科", "—", "—", "78.56%", "-0.10", "70.55%", "-2.14"],
              ["骨科", "—", "—", "63.39%", "-0.11", "67.46%", "-2.53"],
              ["神经外科", "—", "—", "77.49%", "-0.54", "57.77%", "-1.94"],
              ["妇产科学", "65.00%", "-1.06", "65.48%", "-0.75", "66.86%", "-1.33"],
              ["儿科学", "58.71%", "-0.94", "63.95%", "-1.50", "57.93%", "-0.97"],
          ], widths=[2.6, 2.0, 1.7, 2.0, 1.7, 2.0, 1.7], font_size=8.7, first_col_left=True)
add_body(doc, "（1）2025年内科学、外科学、妇产科学和儿科学均低于全国平均水平，其中外科学差距最大，为1.92个百分点；内科学、妇产科学和儿科学分别低于全国1.50、1.33和0.97个百分点。外科学和内科学是临床医学模块近期需要重点加强的学科。")
add_body(doc, "（2）2025年神经内科低于全国4.14个百分点，是内科学细分知识点中差距较大的领域；骨科、肾脏内科、普通外科和神经外科分别低于全国2.53、2.20、2.14和1.94个百分点，应纳入课程组年度重点分析范围。")
add_body(doc, "（3）儿科学原始掌握率在四个临床学科中相对较低，但2025年与全国差距为0.97个百分点，是四个临床学科中差距最小者。后续改进应兼顾基本知识掌握、临床思维和相对全国水平。")
add_body(doc, "（4）妇产科学原始掌握率由2023年的65.00%提高至2025年的66.86%，2025年与全国差距为1.33个百分点。学校应继续巩固课程建设成效，进一步提高病例分析、诊断决策和综合应用能力。")

add_heading_cn(doc, "3.4 医学人文与预防医学模块", 2)
add_caption(doc, "表9  医学人文与预防医学主要学科或知识点与全国差值")
add_table(doc,
          ["学科或知识点", "2023年差值", "2024年差值", "2025年差值"],
          [
              ["卫生法规", "-1.19", "+0.30", "+0.39"],
              ["医学伦理学", "-1.19", "-1.17", "-2.46"],
              ["医学心理学", "-3.44", "-1.25", "-2.33"],
              ["预防医学整体", "-2.19", "-2.35", "-1.89"],
              ["预防医学基本原理与方法", "—", "-2.47", "-3.36"],
          ], widths=[6.0, 3.0, 3.0, 3.0], font_size=10.0, first_col_left=True)
add_body(doc, "卫生法规在2024年和2025年分别高于全国0.30和0.39个百分点，是医学人文模块中相对表现较好的学科。医学心理学连续三年低于全国，2025年差距为2.33个百分点；医学伦理学2025年差距扩大至2.46个百分点。学校应结合2026年医学人文考试大纲调整，完善医学伦理、医学心理、卫生健康法律法规和医学人文实践教学。")
add_body(doc, "预防医学整体连续三年低于全国平均水平，其中基本原理与方法在2024年和2025年分别低于全国2.47和3.36个百分点。应进一步加强流行病学、卫生统计学、疾病预防控制、健康促进和临床预防服务之间的贯通教学，突出医防融合和实际问题解决能力。")

# 四、认知层次
add_heading_cn(doc, "四、认知层次掌握情况", 1)
add_caption(doc, "表10  2025年各模块不同认知层次与全国差值")
add_table(doc,
          ["模块", "记忆层差值", "理解层差值", "应用层差值"],
          [
              ["基础医学", "-0.77", "-1.41", "-3.78"],
              ["医学人文", "-1.67", "-2.17", "-0.79"],
              ["临床医学", "-1.71", "-1.59", "-1.26"],
              ["预防医学", "-1.17", "-2.34", "-1.96"],
          ], widths=[5.0, 3.4, 3.4, 3.4], font_size=10.5)
add_source_note(doc, "注：差值为本校掌握率减全国掌握率，单位为百分点。")
add_body(doc, "2025年各模块记忆、理解和应用层次的掌握率均低于全国平均水平。基础医学应用层差距最大，为3.78个百分点，说明基础知识的情境化应用和综合分析训练需要重点加强；预防医学理解层低于全国2.34个百分点，应进一步强化基本概念、原理和方法的系统理解；医学人文理解层低于全国2.17个百分点，应提高伦理、心理和法律知识在临床情境中的综合运用水平。")
add_body(doc, "学校应在课程考核中合理配置记忆、理解和应用三个层次的试题比例，增加病例型、情境型和综合应用型试题，通过阶段性测验、课程综合考核和临床能力评价持续检验学生知识迁移与问题解决能力。")

# 五、主要问题
add_heading_cn(doc, "五、主要问题", 1)
add_heading_cn(doc, "5.1 总体考试表现的稳定性需要提高", 2)
add_body(doc, "近三年本校总通过率均低于全国平均水平，2024年和2025年百分等级连续为D；医学综合考试通过率在2024年接近全国平均水平，2025年再次回落，平均分与全国差距扩大至9.03分。学校具备接近全国平均水平的基础，但稳定性和相对位次仍需提高。")

add_heading_cn(doc, "5.2 医学综合考试是主要改进环节", 2)
add_body(doc, "三届考生中，约三分之二的最终未通过发生在医学综合考试阶段。学校应将综合知识掌握、临床思维、诊断决策和跨学科整合能力作为下一阶段教学改革与学业支持的重点。")

add_heading_cn(doc, "5.3 基础医学和预防医学存在持续性薄弱环节", 2)
add_body(doc, "基础医学和预防医学模块连续三年低于全国平均水平。解剖学与全国差距连续扩大，医学免疫学在2025年出现明显下降，预防医学基本原理与方法差距扩大。基础医学应用层是2025年认知层次中差距最大的领域，说明基础知识与临床问题解决之间的衔接仍需加强。")

add_heading_cn(doc, "5.4 临床医学部分学科和知识点需要重点加强", 2)
add_body(doc, "2025年外科学、内科学、妇产科学和儿科学均低于全国平均水平，神经内科、骨科、肾脏内科、普通外科和神经外科等知识点差距相对较大。学校应根据课程目标、校内考核、学生反馈和临床教学实际，逐项分析教学内容、师资、实践机会和考核方式，明确改进重点。")

add_heading_cn(doc, "5.5 学业预警和毕业后支持需要进一步完善", 2)
add_body(doc, "2025年全日制应届本科考生医学综合考试最低分为100分，存在需要重点关注的个案。学校应完善分数段统计，重点关注300分以下、300—329分和330—359分考生，结合课程不及格、补考、实习表现和阶段性综合测试结果，提前识别学习困难学生并开展针对性指导。")
add_body(doc, "2023—2025年非应届考生医学综合考试通过率分别为38.98%、50.20%和45.95%。非应届考生数据应单独分析，学校应完善毕业生联系、政策咨询、学习资源和备考指导服务。")

# 六、标准对照
add_heading_cn(doc, "六、与《中国本科医学教育标准——临床医学专业（2022版）》相关要求的对照分析", 1)
add_body(doc, "医师资格考试成绩是评价人才培养结果的重要外部指标。本节依据考试结果梳理相关标准领域的重点工作，正式认证结论由认证程序依据人才培养全过程材料综合形成。")
add_caption(doc, "表11  相关标准领域的重点工作")
add_table(doc,
          ["标准领域", "考试结果反映的重点", "下一阶段重点工作"],
          [
              ["课程计划", "基础医学、预防医学连续低于全国，部分临床知识点差距扩大", "完善培养目标、毕业要求、课程目标、教学内容和考核要求之间的对应关系，加强基础、临床、预防和人文课程整合"],
              ["学业考核与评价", "医学综合考试年度表现波动，基础医学应用层差距较大", "完善形成性评价、终结性评价和临床能力评价，增加病例型、情境型和综合应用型考核"],
              ["学生", "存在低分个案，非应届考生通过率相对较低", "健全早期预警、分类指导、心理与学习咨询以及毕业后支持机制"],
              ["教师", "薄弱学科和基础临床协同教学需要加强", "优化师资结构，推进联合备课、协同授课和医学教育能力培训"],
              ["教育资源", "实践技能通过率连续两年百分等级为D", "提高技能中心、模拟教学和临床教学基地资源的开放程度与使用效益"],
              ["教育评价", "年度成绩分析向改进措施转化仍需加强", "明确责任部门、改进目标、完成时限、过程指标和效果评价，形成持续改进机制"],
          ], widths=[2.5, 5.4, 7.1], font_size=9.0, first_col_left=True)
add_body(doc, "课程计划方面，应围绕毕业要求，系统梳理基础医学、临床医学、预防医学、医学人文和中医学基础课程的知识、能力和职业素养要求，重点加强基础与临床、临床与预防、专业教育与人文教育的衔接。")
add_body(doc, "学业考核与评价方面，应建立课程考核、阶段性综合测试、OSCE、Mini-CEX、DOPS、病例讨论和临床推理评价相结合的多元评价体系，强化形成性反馈和综合应用能力考查。")
add_body(doc, "教师与教育资源方面，应加强基础医学教师和临床教师协作，完善临床教学基地统一标准、教师培训、学生评价和年度质量评估，提高临床技能中心、附属医院和教学医院对本科教学的支持能力。")
add_body(doc, "教育评价方面，应将国家医师资格考试结果与校内课程评价、毕业要求达成评价、临床能力评价和毕业生发展情况统筹分析，按年度研究改进措施并评价实施效果。")

# 七、整改建议
add_heading_cn(doc, "七、整改建议", 1)
add_heading_cn(doc, "（一）完善医师资格考试分析与持续改进机制", 2)
add_item(doc, 1, "健全校级协调机制。", "由医学部牵头，教务处、教学质量监控与评估处、基础医学院、临床医学院、附属医院和相关教学单位共同参与，统一数据口径，明确职责分工，制定年度工作安排，将考试分析纳入临床医学专业建设和人才培养质量评价。", "医学部")
add_item(doc, 2, "完善年度成绩分析。", "每年从实践技能和医学综合考试通过率、平均分、分数段、模块、专业知识点、考核要点和认知层次等方面开展分析；持续性薄弱环节与年度异常变化分别研究。分析工作同步参考课程考核、试题质量、师资与教学基地情况、学生和教师反馈，形成年度专题报告。", "教学质量监控与评估处、医学部")
add_item(doc, 3, "落实改进责任和报告制度。", "针对年度分析发现的问题，形成书面整改方案，逐项明确责任单位、改进目标、具体措施、完成时限、过程指标、结果指标和验收材料。年度进展提交学校教学指导委员会审议，重大事项按程序提交校长办公会研究。", "医学部、教学质量监控与评估处")

add_heading_cn(doc, "（二）加强课程教学与考核评价衔接", 2)
add_item(doc, 4, "梳理课程教学与培养要求的对应关系。", "坚持人才培养目标和毕业要求的主导地位，系统梳理毕业要求、课程目标、教学内容、教学方法、校内考核与行业准入要求之间的对应关系，重点检查课程覆盖、内容重复、衔接不足和综合应用要求落实情况。", "教务处、医学部")
add_item(doc, 5, "深化校内考核改革。", "依据课程性质合理配置形成性评价和终结性评价权重，完善试卷命题方案，增加病例型、情境型和综合应用型试题，开展试题难度、区分度、信度和选项功能分析。课程总评应综合反映知识、能力和职业素养目标的达成情况。", "教务处、基础医学院、临床医学院")
add_item(doc, 6, "开展分阶段综合测试。", "在基础医学阶段、临床课程阶段和实习阶段设置综合测试，内容覆盖基础、临床、预防和人文领域，形成学生个体和年级学习发展资料，用于课程改进、学业指导和毕业要求达成评价。", "医学部、教务处")

add_heading_cn(doc, "（三）强化薄弱环节整改与基础临床融合", 2)
add_item(doc, 7, "开展重点学科专项改进。", "将解剖学、预防医学基本原理与方法、基础医学应用层列为持续改进重点，将医学免疫学、神经内科、骨科、肾脏内科、普通外科等列入2025年度重点分析范围。相关课程组应检查课程目标、教学内容、学时结构、师资稳定性、学生学习投入和校内考核结果，提出年度改进措施。", "基础医学院、临床医学院")
add_item(doc, 8, "推进基础医学与临床医学协同教学。", "建立基础医学教师与临床教师联合备课、案例共建和协同授课机制，将解剖、生理、病理、免疫、药理等基础知识融入临床问题解决情境，通过病例分析、综合应用题和分阶段综合测试评价教学效果。", "医学部、基础医学院、临床医学院")
add_item(doc, 9, "加强预防医学和医学人文教学。", "围绕医防融合、健康促进、临床预防服务、沟通、人文关怀、伦理与法律责任，将相关内容贯穿基础课程、临床技能训练和实习教学。依据正式发布的考试大纲和课程标准及时更新教学内容，完善情境模拟、伦理案例讨论和卫生健康法律法规实务训练。", "基础医学院、临床医学院、护理学院")

add_heading_cn(doc, "（四）加强临床实践教学与教学基地质量管理", 2)
add_item(doc, 10, "完善临床技能训练和能力评价。", "建立覆盖基本技能、综合技能、临床思维和医患沟通的递进式训练体系，分阶段实施OSCE、Mini-CEX、DOPS、病例讨论和临床推理评价，统一考站设置、考核内容、评分标准、考官培训和质量控制。技能中心开放安排应结合课程计划和学生训练需求确定。", "临床医学院、附属医院")
add_item(doc, 11, "加强教学基地年度质量评估。", "统一附属医院和教学医院的临床教学标准，按基地分析学生课程成绩、临床能力评价、轮转完成情况、师资配置和学生评价，形成教学基地年度质量分析报告。对差距较大的基地开展专项改进和复查，持续提高临床教学质量的一致性。", "医学部、临床医学院、附属医院")

add_heading_cn(doc, "（五）完善学生学业支持与毕业后跟踪服务", 2)
add_item(doc, 12, "健全学业预警和分类指导。", "重点统计300分以下、300—329分、330—359分等考生群体，结合课程不及格、补考、分阶段综合测试和实习表现，提前识别学习困难学生，制定个人学习支持方案。过程评价重点关注预警覆盖率、指导完成率和后续学习改善情况。", "基础医学院、临床医学院、学生工作部门")
add_item(doc, 13, "完善毕业生跟踪服务。", "对毕业后未通过医师资格考试的校友，建立常态化联系机制，定期提供政策解读、考试大纲、学习资源、线上辅导和经验交流。非应届考生单独统计和分析，相关结果用于改进毕业后支持服务。", "临床医学院")
add_item(doc, 14, "完善工作评价方式。", "教师教学质量评价实行团队评价和多指标综合评价，重点考察课程建设、教学实施、学生能力达成和持续改进成效；学生支持工作以学习发展、能力提升和服务效果为主要评价内容。涉及学生成绩告知和家校沟通的事项，按照学校制度和学生知情原则规范实施。", "教务处、医学部、学生工作部门")

# 八、目标
add_heading_cn(doc, "八、未来三年工作目标", 1)
add_body(doc, "未来三年以相对全国表现、连续年度结果和关键过程指标为重点，实行年度评价和适时调整。")
add_caption(doc, "表12  未来三年主要目标")
add_table(doc,
          ["指标", "2025年情况", "三年目标"],
          [
              ["总通过率与全国差距", "低于全国2.38个百分点", "控制在±1个百分点以内"],
              ["总通过率连续性", "2024年和2025年连续低于全国，均为D级", "连续两年达到或超过全国平均水平，百分等级达到C级及以上"],
              ["医学综合考试平均分差距", "低于全国9.03分", "控制在3分以内"],
              ["实践技能考试表现", "低于全国1.33个百分点，D级", "稳定达到或超过全国平均水平，百分等级达到C级及以上"],
              ["重点知识点相对差距", "解剖学-5.26个百分点；医学免疫学-5.63个百分点", "与全国差距较2025年缩小50%以上"],
              ["核心课程对应关系梳理", "尚未形成统一校级成果", "核心课程全部完成梳理并通过审核"],
              ["分阶段综合测试和试题分析", "尚未形成完整体系", "覆盖主要培养阶段；核心课程年度试题分析率100%"],
              ["教学基地质量评估和学生指导", "缺少统一结果指标", "教学基地年度质量评估率100%；重点预警学生指导完成率不低于90%"],
          ], widths=[4.1, 5.5, 6.0], font_size=9.0, first_col_left=True)
add_body(doc, "百分等级由当年参评学校的相对位置确定，固定通过率与百分等级分别评价。学校每年根据考试政策、试题结构、参评范围和实际工作进展，对目标完成情况进行评估；目标调整事项提交学校教学指导委员会审议。")

# 九、结语
add_heading_cn(doc, "九、结语", 1)
add_body(doc, "2023—2025年，九江学院全日制应届本科考生临床执业医师资格考试总通过率总体接近全国水平，但连续三年低于全国平均，2024年和2025年百分等级均为D，年度表现和相对位次需要进一步提高。分阶段看，主要改进空间集中在医学综合考试。基础医学和预防医学连续三年低于全国，解剖学为持续性相对薄弱学科，医学免疫学在2025年出现明显下降；基础医学应用层是2025年认知层次中与全国差距最大的领域。")
add_body(doc, "学校应坚持人才培养目标和毕业要求的主导地位，以临床医学专业认证和持续改进为主线，重点加强考试成绩分析、课程教学与考核评价衔接、基础与临床教学融合、临床实践教学与教学基地质量管理以及学生学业支持。各责任部门应按照年度整改方案推进工作，学校教学指导委员会定期审议进展，校长办公会研究解决重大事项，持续提高临床医学人才培养质量和毕业生岗位胜任力。")

# 十、附件
add_heading_cn(doc, "十、附件：数据口径与分析方法说明", 1)
add_heading_cn(doc, "1. 分析对象", 2)
add_body(doc, "本报告以全日制应届本科考生为主要分析对象。国家医学考试中心所称应届本科考生，是指毕业后一年参加考试的本科考生。2023年、2024年和2025年分别对应2022届、2023届和2024届毕业生。全体考生包括应届、非应届及不同学历层次考生，其数据用于呈现学校考试总体情况和毕业后支持需求。")

add_heading_cn(doc, "2. 数据来源与内容范围", 2)
add_body(doc, "数据来源为国家医学考试中心提供的2023—2025年《医师资格考试成绩分析报告（临床执业医师）》。报告数据以考生报名信息为基础，医学综合考试成绩包括当年“一试”，未包括“二试”；医学综合考试通过人数按四个单元成绩统计，未包括军事医学、院前急救和儿科加试成绩。")

add_heading_cn(doc, "3. 通过率计算口径", 2)
add_body(doc, "实践技能考试通过率＝实践技能考试通过人数÷实践技能考试实考人数。医学综合考试通过率＝医学综合考试通过人数÷医学综合考试实考人数。总通过率＝医学综合考试通过人数÷（实践技能考试实考人数＋实践技能考试免考人数）。三类指标分母不同，分别反映实践技能、医学综合考试和完整考试过程的通过情况。")

add_heading_cn(doc, "4. 百分等级", 2)
add_body(doc, "国家医学考试中心将参评学校某项通过率由高到低排列，划分为A、B、C、D、E五个等级，各等级学校数量约占20%。A为前20%，B为20%—40%，C为40%—60%，D为60%—80%，E为后20%。2023年学校报告未提供百分等级，2024年和2025年报告提供了相关等级。")

add_heading_cn(doc, "5. 知识分类与年度比较", 2)
add_body(doc, "2023年报告主要按学科呈现掌握率，2024年和2025年进一步按专业知识点、系统、考核要点和认知层次呈现，并新增中医学基础模块。细分知识点的连续比较主要采用2024—2025年数据。2023—2025年国家报告均使用“解剖学”名称，本报告统一采用该名称。")
add_body(doc, "国家医学考试中心未提供跨年度试题等值量表。年度比较以同期全国差值、百分等级、平均分和连续年度趋势为主要依据，课程和知识点改进效果同时参考校内考核、试题质量、教学实施和学生学习情况。")

add_heading_cn(doc, "6. 原因分析与改进依据", 2)
add_body(doc, "考试结果用于确定需要重点关注的模块、学科、知识点和能力层次。具体原因由课程目标与内容、教学方法、试题质量、师资结构、教学基地、学生学习投入和反馈等资料综合确认。整改措施应设置过程指标和结果指标，并通过后续课程评价、临床能力评价和医师资格考试结果检验实施成效。")

add_heading_cn(doc, "7. 补充数据说明", 2)
add_body(doc, "2023—2025年非应届考生医学综合考试通过率分别为38.98%、50.20%和45.95%。2025年全日制应届本科考生医学综合考试最低分为100分，学校应进一步完善分数段统计和个案分析。医学综合考试合格分数线为360分，实践技能考试合格分数线为60分。")

add_heading_cn(doc, "主要资料来源", 2)
for text in [
    "1. 国家医学考试中心：《2023年医师资格考试成绩分析报告（九江学院，临床执业医师）》。",
    "2. 国家医学考试中心：《2024年医师资格考试成绩分析报告（九江学院，临床执业医师）》。",
    "3. 国家医学考试中心：《2025年医师资格考试成绩分析报告（九江学院，临床执业医师）》。",
    "4. 教育部临床医学专业认证工作委员会：《中国本科医学教育标准——临床医学专业（2022版）》。",
]:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(26)
    r = p.add_run(text)
    set_run_font(r, chinese=BODY_FONT, size=15)

# Quality checks for wording explicitly rejected by the user.
text = all_document_text(doc)
banned_terms = [
    "台账", "蓝图", "画像", "数据治理", "年度诊断", "根因", "复盘", "精准支持",
    "尚不能", "不能据此", "不能仅据", "不宜据此", "不宜仅凭", "据此认定",
    "相对教学成效持续改善", "仅凭原始掌握率", "风险提示"
]
found = [term for term in banned_terms if term in text]
if found:
    raise RuntimeError("发现未清理表述：" + "、".join(found))

# Keep table headers and headings together where possible.
for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        p.paragraph_format.keep_with_next = True

# Save.
doc.save(OUTPUT)
print(f"Generated: {OUTPUT}")
print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
